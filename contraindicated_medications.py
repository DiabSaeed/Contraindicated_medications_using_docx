from docx import Document
from typing import Dict, List, Tuple

class ContraIndicated:
    doc = Document()
    headers = ['Medication Name', 'Species', 'Condition', 'Notes']

    def __init__(self, sp_name: str, med_name: str, med_condition: str, med_notes: str):
        """
        Initialize a new medication instance with species and medication details.
        """
        self.sp_name = sp_name
        self.med_name = med_name
        self.med_condition = med_condition
        self.med_notes = med_notes
    @staticmethod
    def add_header(header):
        ContraIndicated.doc.add_heading(header, 0)
    def add_medication(self, species_data: Dict[str, List[Tuple[str, str, str, str]]]) -> Dict[str, List[Tuple[str, str, str, str]]]:
        """
        Add a medication for the given species to the species_data dictionary.
        """
        if self.sp_name not in species_data:
            species_data[self.sp_name] = []
        species_data[self.sp_name].append((self.med_name, self.sp_name, self.med_condition, self.med_notes))
        return species_data

    @staticmethod
    def create_table_for_species(species_data: Dict[str, List[Tuple[str, str, str, str]]]):
        """
        Generate a detailed table for each species with medications, conditions, and notes.
        """
        for species, data in species_data.items():
            ContraIndicated.doc.add_heading(f'{species}', level=1)
            table = ContraIndicated.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(ContraIndicated.headers):
                hdr_cells[i].text = header

            
            for medication, species, condition, notes in data:
                row_cells = table.add_row().cells
                row_cells[0].text = medication
                row_cells[1].text = species
                row_cells[2].text = condition
                row_cells[3].text = notes

    @staticmethod
    def add_to_doc(file_path: str):
        """
        Save the document to the given file path.
        """
        ContraIndicated.doc.save(file_path)



species_data_detailed = {
    "Cats": [
        ("Acetaminophen", "Cats", "", "Highly toxic, can cause liver failure and methemoglobinemia"),
        ("Aspirin", "Cats", "Long-term use", "Slow metabolism, can lead to toxicity with prolonged use"),
        ("Ivermectin", "Cats", "", "Toxic in high doses, can cause neurological issues like seizures"),
        ("Enrofloxacin", "Cats", "High doses", "Retinal toxicity, may cause blindness"),
        ("Cisapride", "Cats", "Pre-existing cardiac conditions", "Can cause serious arrhythmias"),
        ("Methimazole", "Cats", "Liver disease", "Can exacerbate hepatic issues"),
    ],
    "Dogs": [
        ("Ibuprofen", "Dogs", "", "Causes stomach ulcers, kidney damage, and can be fatal"),
        ("Dexamethasone", "Dogs", "Pregnancy", "Risk of abortion or fetal malformation"),
        ("Xylitol", "Dogs", "", "Highly toxic, even small amounts cause hypoglycemia and liver failure"),
        ("Ivermectin", "Dogs", "Collie breeds (MDR1 mutation)", "Neurological toxicity in susceptible breeds"),
        ("Tetracycline", "Dogs", "Young growing dogs", "Can cause enamel hypoplasia and tooth discoloration"),
        ("Acepromazine", "Dogs", "Heart disease", "Can cause hypotension"),
    ],
    "Horses": [
        ("Flunixin meglumine", "Horses", "Renal insufficiency", "Can worsen kidney function"),
        ("Dexamethasone", "Horses", "Laminitis-prone horses", "Increases the risk of laminitis"),
        ("Phenylbutazone", "Horses", "Prolonged use", "Gastrointestinal ulceration and kidney damage"),
        ("Ivermectin", "Horses", "Foals and small ponies", "Toxicity can occur due to overdosing"),
        ("Atropine", "Horses", "Colic or ileus", "Can exacerbate gut motility issues"),
        ("Xylazine", "Horses", "Respiratory disease", "May cause respiratory depression"),
    ],
    "Cattle": [
        ("Chloramphenicol", "Cattle", "", "Prohibited in food animals due to human health risks (aplastic anemia)"),
        ("Dexamethasone", "Cattle", "Pregnancy", "Can cause abortion or fetal malformations"),
        ("Tilmicosin", "Cattle", "", "Fatal if injected intravenously; causes cardiac failure"),
        ("Ivermectin", "Cattle", "", "Can be toxic in overdoses, especially in young calves"),
        ("Gentamicin", "Cattle", "", "Prohibited in food animals due to residue concerns"),
        ("Clenbuterol", "Cattle", "", "Prohibited in food-producing animals due to potential human toxicity"),
    ],
}

medication_one = ContraIndicated("Cats", "Acetaminophen", "", "Deficiency of enzyme")
medication_one.add_header('Detailed Guide: Contraindicated Medications for Various Species')
species_data_detailed = medication_one.add_medication(species_data_detailed)
ContraIndicated.create_table_for_species(species_data_detailed)
ContraIndicated.add_to_doc("E:/programming/Detailed_Contraindicated_Medications_Veterinary.docx")
